from openpyxl import load_workbook
import os
from docx import Document
from docx2pdf import convert
from RPA.Email.ImapSmtp import ImapSmtp
from dotenv import load_dotenv
from docx.shared import Pt

class Excel:
    def __init__(self, file_path):
        self.file_path = file_path
        self.wb = load_workbook(self.file_path)
        self.sheet = self.wb.active
    
    def get_total_rows(self):
        return self.sheet.max_row - 1
    
    def get_headers(self):
        if self.get_total_rows() > 0:
            return self.read_from_file(1)

    def read_from_file(self, row_to_read):
        for row in self.sheet.iter_rows(min_row=row_to_read, max_row=row_to_read, values_only=True):
            return row

class Invoice:
    def __init__(self, invoice_template, invoice_folder):
        self.invoice_template = os.path.join(".", invoice_template)
        self.invoice_folder = os.path.join(".", invoice_folder)
    
    def generate_invoices(self, invoice_num, data):
        data["invoice_number"] = invoice_num
        template = Document(self.invoice_template)

        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text = "".join(run.text for run in para.runs).strip()
                        new_text = full_text
                        for key, value in data.items():
                            new_text = new_text.replace(f"[{key}]", str(value))
                        if new_text != full_text:
                            para.clear()
                            run = para.add_run(new_text)
                            if "[company_name]" in full_text:
                                run.bold = True
                                run.font.size = Pt(20)
                            elif full_text in ["[client_name]","[invoice_number]", "[shipped_to_client]"]:
                                run.bold = True

        
        for table in template.tables:
            if len(table.columns) == 4 and table.cell(0, 0).text.strip().upper() == "DESCRIPTION":
                if len(table.rows) > 1:
                    row = 1
                    for item in data["items"]:
                        table.add_row()
                        new_row = table.rows[-1]
                        current_tr = table.rows[row - 1]._tr
                        moving_tr = new_row._tr
                        table._tbl.remove(moving_tr)
                        current_tr.addnext(moving_tr)
                        inserted_row = table.rows[row]
                        inserted_row.cells[0].text = item["product"]
                        inserted_row.cells[1].text = str(item["quantity"])
                        inserted_row.cells[2].text = str(item["rate"])
                        inserted_row.cells[3].text = str(item["quantity"] * item["rate"])

                        row += 1
                row_to_remove = table.rows[row]._tr
                table._tbl.remove(row_to_remove)
                break
        
        for table in template.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip() == "SUBTOTAL":
                            next_cell = row.cells[-1]
                            next_cell.text = str(data["subtotal"])
                        elif paragraph.text.strip() == "TAX":
                            next_cell = row.cells[-1]
                            next_cell.text = str(data["total_tax"])
                        elif paragraph.text.strip() == "SHIPPING & HANDLING":
                            next_cell = row.cells[-1]
                            next_cell.text = str(data["shipping_charges"])
                        elif paragraph.text.strip() == "DISCOUNT":
                            next_cell = row.cells[-1]
                            next_cell.text = str(data["discount"])
                        elif paragraph.text.strip() == "GRAND TOTAL":
                            next_cell = row.cells[-1]
                            next_cell.text = str(data["total"])

            
        # Save the edited document
        # check if invoice is good after multiple elements
        file_name = os.path.join(self.invoice_folder, data["invoice_number"])
        docx_path = f"{file_name}.docx"
        pdf_path = f"{file_name}.pdf"

        template.save(docx_path)
        convert(docx_path, pdf_path)
        os.remove(docx_path)
        return pdf_path
    
class DataFormat:
    def __init__(self, company_name, address, contact, tax_percent, data_headers,min_discount_amount,discount_percent):
        self.company_name = company_name
        self.address = address
        self.contact = contact
        self.tax = tax_percent / 100
        self.min_discount_amount = min_discount_amount
        self.discount = discount_percent/100
        self.invoice_list = set()
        self.invoices = {}
    
    def read_invoices(self, details):
        invoice_id = details[0]

        item_subtotal = int(details[6]) * int(details[7])
        item_tax = item_subtotal * self.tax
        shipping = int(details[11])

        if invoice_id not in self.invoice_list:
            data = {
                "company_name": self.company_name,
                "address" : self.address,
                "contact": self.contact,
                "client_name": details[1],
                "shipped_to_client": details[12],
                "client_email": details[2],
                "invoice_date": details[3].strftime('%d-%m-%Y'),
                "due_date": details[4].strftime('%d-%m-%Y'),
                "items": [{
                    "product": details[5],
                    "quantity": details[6],
                    "rate": details[7]
                }],
                "payment_mode": details[8],
                "billing_address": details[9],
                "shipping_address": details[10],
                "shipping_charges": shipping,
                "subtotal": item_subtotal,
                "total_tax": item_tax
            }

            self.invoices[invoice_id] = data
            self.invoice_list.add(invoice_id)
        else:
            self.invoices[invoice_id]["items"].append({
                "product": details[5],
                "quantity": details[6],
                "rate": details[7]
            })
            self.invoices[invoice_id]["subtotal"] += item_subtotal
            self.invoices[invoice_id]["total_tax"] += item_tax
            self.invoices[invoice_id]["shipping_charges"] += shipping

        invoice = self.invoices[invoice_id]
        total = invoice["subtotal"] + invoice["total_tax"] + invoice["shipping_charges"]
        discount = 0
        if total >= self.min_discount_amount:
            discount = total * self.discount
            total -= discount
        invoice["discount"] = discount
        invoice["total"] = total


class Email:
    def __init__(self):
        self.mail = ImapSmtp()
        self.gmail_account = os.getenv("GMAIL_ACCOUNT")
        self.gmail_password = os.getenv("GMAIL_PASSWORD")

        self.mail.authorize(
            account=self.gmail_account,
            password=self.gmail_password,
            smtp_server="smtp.gmail.com",
            smtp_port=587,
        )

    def send_mail(self, file_path, invoice_num, client_name, client_email=None):
        if not client_email:
            client_email = self.gmail_account
        if not client_email:
            raise ValueError("No valid recipient email provided.")

        body = f"Dear {client_name}, \nPlease find attached the invoice {invoice_num} for the recent purchase/service. \nIf you have any questions or concerns, feel free to contact us.\nThank you for your business!\nBest regards,\nAnushka\n"

        self.mail.send_message(
            recipients=client_email,
            cc=self.gmail_account,
            sender=self.gmail_account,
            subject=f"Invoice {invoice_num} from Anushka",
            body=body,
            attachments=file_path
        )

class InvoiceGenerator:
    def __init__(self, company_name, address, contact, file_name, tax_percent, invoice_template, invoice_folder,min_discount_amount,discount_percent):
        self.company_name = company_name
        self.address = address
        self.contact = contact
        self.file_path = os.path.join(".", file_name)
        self.tax = tax_percent
        self.invoice_template = invoice_template
        self.invoice_folder = invoice_folder
        self.min_discount_amount = min_discount_amount
        self.discount_percent = discount_percent

    def main(self):
        load_dotenv()
        excel_work = Excel(self.file_path)
        headers = excel_work.get_headers()
        
        invoice_generator = Invoice(self.invoice_template, self.invoice_folder)
        email = Email()
        
        data_formatter = DataFormat(self.company_name, self.address, self.contact, self.tax, headers,self.min_discount_amount, self.discount_percent)

        total_invoices = excel_work.get_total_rows()

        # for i in range(2, 10):
        for i in range(2, total_invoices + 2):
            details = excel_work.read_from_file(i)
            data_formatter.read_invoices(details)
        
        for invoice_num,data in data_formatter.invoices.items():
            file_path = invoice_generator.generate_invoices(invoice_num, data)
            email.send_mail(file_path, invoice_num, data["client_name"], data["client_email"])
            print(file_path)

if __name__ == "__main__":
    company_name = input("Enter name of your company: ")
    line1 = input("Enter address line 1: ")
    line2 = input("Enter address line 2: ")
    address = f"{line1} \n{line2}"
    contact = input("Enter contact number: ")
    invoice = InvoiceGenerator(company_name, address, contact, "invoice_details.xlsx", 10, "invoice-basic.docx", "invoices", 2000,5)
    invoice.main()